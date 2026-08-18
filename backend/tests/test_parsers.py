import io

import docx
import pytest
from app.rag.parsers import parse_document
from pypdf import PdfWriter


def _make_pdf() -> bytes:
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def test_parse_txt_single_page():
    pages = parse_document("你好\n世界\n".encode(), "txt")
    assert len(pages) == 1
    assert "你好" in pages[0].text
    assert pages[0].page_no == 0


def test_parse_md_single_page():
    pages = parse_document("# 标题\n正文".encode(), "md")
    assert len(pages) == 1
    assert pages[0].text.startswith("# 标题")


def test_parse_docx_paragraphs():
    d = docx.Document()
    d.add_paragraph("第一段")
    d.add_paragraph("第二段")
    buf = io.BytesIO()
    d.save(buf)
    pages = parse_document(buf.getvalue(), "docx")
    assert "第一段" in pages[0].text and "第二段" in pages[0].text


def test_parse_pdf_returns_one_page():
    pages = parse_document(_make_pdf(), "pdf")
    assert len(pages) == 1
    assert pages[0].page_no == 0


def test_unsupported_type_raises():
    with pytest.raises(ValueError):
        parse_document(b"x", "exe")
